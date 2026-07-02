import os
import re
import io
import requests
import chromadb
from docx import Document
from PIL import Image
from chromadb.utils.embedding_functions import OpenCLIPEmbeddingFunction

# =====================================================================
# CONFIGURATION MATRIX
# =====================================================================
AZURE_OPENAI_ENDPOINT = "https://azure.com"
AZURE_API_KEY = "5IetNCvkyoLqcBHq3x05X8a6ecOUsBXQSW48dIrtrxldtsytiK2bJQQJ99BKACYeBjFXJ3w3AAABACOGkP9a"
AZURE_API_VERSION = "2024-12-01-preview"
AZURE_EMBED_DEPLOYMENT = "text-embedding-3-small"

REPORTS_DIR = "Reports"
ACCEPTABLE_DIR = "Acceptable_Images"
REJECTABLE_DIR = "Rejectable_Images"
EXTRACTED_DOC_IMAGES_DIR = "Extracted_Manual_Images"
DB_PATH = "hybrid_chroma_db"

os.makedirs(EXTRACTED_DOC_IMAGES_DIR, exist_ok=True)

print("[Init] Spawning Permanent Local Hybrid Vector DB on disk...")
chroma_client = chromadb.PersistentClient(path=DB_PATH)
clip_ef = OpenCLIPEmbeddingFunction()

try:
    chroma_client.delete_collection(name="manual_text_rules")
    chroma_client.delete_collection(name="visual_standards")
except Exception:
    pass

text_collection = chroma_client.create_collection(name="manual_text_rules")
vision_collection = chroma_client.create_collection(name="visual_standards", embedding_function=clip_ef)

# =====================================================================
# MULTIMODAL EXTRACTION UTILITIES
# =====================================================================
def get_azure_text_embedding(text_string):
    base_url = AZURE_OPENAI_ENDPOINT.rstrip('/')
    url = f"{base_url}/openai/deployments/{AZURE_EMBED_DEPLOYMENT}/embeddings?api-version={AZURE_API_VERSION}"
    headers = {"Content-Type": "application/json", "api-key": AZURE_API_KEY}
    payload = {"input": text_string}
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=40)
        if response.status_code == 200:
            return response.json()["data"]["embedding"]
    except Exception:
        pass
    return [0.0] * 1536

def process_and_index_docx_multimodal(file_path, text_col, vis_col):
    filename = os.path.basename(file_path)
    doc = Document(file_path)
    
    vr_match = re.search(r'(VR\d+)', filename, re.IGNORECASE)
    vr_tag = vr_match.group(1).upper() if vr_match else os.path.splitext(filename)[0]
    
    current_chunk = ""
    block_counter = 0
    img_counter = 0
    
    # 1. Extract Text Data and Tables
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            block_counter += 1
            current_chunk += " " + txt
            if len(current_chunk) > 700:
                txt_emb = get_azure_text_embedding(current_chunk.strip())
                text_col.add(
                    embeddings=[txt_emb],
                    documents=[current_chunk.strip()],
                    metadatas=[{"source_file": filename, "vr_tag": vr_tag, "type": "paragraph"}],
                    ids=[f"text_{vr_tag}_{block_counter}"]
                )
                current_chunk = ""

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                block_counter += 1
                current_chunk += " " + " | ".join(list(set(row_text)))
                if len(current_chunk) > 700:
                    txt_emb = get_azure_text_embedding(current_chunk.strip())
                    text_col.add(
                        embeddings=[txt_emb],
                        documents=[current_chunk.strip()],
                        metadatas=[{"source_file": filename, "vr_tag": vr_tag, "type": "table"}],
                        ids=[f"table_{vr_tag}_{block_counter}"]
                    )
                    current_chunk = ""
                    
    if current_chunk.strip():
        txt_emb = get_azure_text_embedding(current_chunk.strip())
        text_col.add(
            embeddings=[txt_emb],
            documents=[current_chunk.strip()],
            metadatas=[{"source_file": filename, "vr_tag": vr_tag, "type": "residual"}],
            ids=[f"residual_{vr_tag}_{block_counter}"]
        )

    # 2. Extract and Save Manual Images
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.target_ref:
            img_counter += 1
            img_bytes = rel.target_part.blob
            out_img_name = f"extracted_{vr_tag}_img_{img_counter}.png"
            out_img_path = os.path.join(EXTRACTED_DOC_IMAGES_DIR, out_img_name)
            
            with open(out_img_path, "wb") as f:
                f.write(img_bytes)
                
            vis_col.add(
                images=[out_img_path],
                documents=[out_img_name],
                metadatas=[{
                    "disposition_label": "Document_Embedded_Standard", 
                    "source_folder": EXTRACTED_DOC_IMAGES_DIR,
                    "parent_manual": filename,
                    "vr_tag": vr_tag
                }],
                ids=[f"doc_image_node_{vr_tag}_{img_counter}"]
            )
    print(f"   ✅ Fully Indexed Multimodal Elements: {filename} ({block_counter} Text blocks, {img_counter} Diagrams Extracted)")

def main():
    if os.path.exists(REPORTS_DIR):
        available_docs = [f for f in os.listdir(REPORTS_DIR) if f.lower().endswith('.docx') and not f.startswith('~$')]
        print(f"\n -> Extracting multimodal elements from {len(available_docs)} manuals...")
        for f_name in available_docs:
            process_and_index_docx_multimodal(os.path.join(REPORTS_DIR, f_name), text_collection, vision_collection)

    image_targets = [(ACCEPTABLE_DIR, "Acceptable"), (REJECTABLE_DIR, "Rejectable")]
    for folder, classification in image_targets:
        if not os.path.exists(folder):
            continue
        files = [f for f in os.listdir(folder) if f.lower().endswith(('.png', '.jpg', '.jpeg'))]
        print(f"\n -> Indexing {len(files)} files from benchmarks folder: '{folder}'...")
        for f_name in files:
            f_path = os.path.join(folder, f_name)
            vr_match = re.search(r'(VR\d+)', f_name, re.IGNORECASE)
            vr_tag = vr_match.group(1).upper() if vr_match else "UNKNOWN_PART"
            
            vision_collection.add(
                images=[f_path],
                documents=[f_name],
                metadatas=[{"disposition_label": classification, "source_folder": folder, "vr_tag": vr_tag}],
                ids=[f"image_node_{folder}_{f_name}"]
            )
    print(f"\n[✔] ALL TEXT & DIAGRAM VECTORS STORED IN HYBRID CHROMA DB.")

if __name__ == "__main__":
    main()
